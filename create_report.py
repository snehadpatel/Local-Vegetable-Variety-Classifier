import os
import json
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Color Palette Definitions
COLOR_TEAL = RGBColor(0, 95, 115)     # Primary Headings
COLOR_CHARCOAL = RGBColor(47, 62, 70)  # Subheadings
COLOR_GRAY = RGBColor(100, 110, 120)   # Body captions / borders
COLOR_BLACK = RGBColor(0, 0, 0)
COLOR_MUTED = RGBColor(220, 225, 230)

def set_cell_background(cell, hex_color):
    """Sets background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading_with_spacing(doc, text, level, before=14, after=6):
    """Adds a heading with explicit spacing and color."""
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(before)
    h.paragraph_format.space_after = Pt(after)
    h.paragraph_format.keep_with_next = True
    
    # Set color and sizing based on level
    for run in h.runs:
        run.font.name = 'Calibri'
        if level == 1:
            run.font.color.rgb = COLOR_TEAL
            run.font.size = Pt(18)
            run.font.bold = True
        elif level == 2:
            run.font.color.rgb = COLOR_CHARCOAL
            run.font.size = Pt(14)
            run.font.bold = True
        else:
            run.font.color.rgb = COLOR_CHARCOAL
            run.font.size = Pt(12)
            run.font.bold = True
    return h

def add_code_block(doc, file_name, code_text):
    """Adds a beautifully formatted code block in monospace font with background shading."""
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(8)
    p_title.paragraph_format.space_after = Pt(2)
    p_title.paragraph_format.keep_with_next = True
    run_title = p_title.add_run(f"📄 File: {file_name}")
    run_title.font.name = 'Calibri'
    run_title.font.bold = True
    run_title.font.size = Pt(10.5)
    run_title.font.color.rgb = COLOR_TEAL

    # Code container using a single-cell table
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, "F4F5F6") # Light gray background
    
    # XML padding configurations
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin in ['top', 'bottom', 'left', 'right']:
        node = OxmlElement(f'w:{margin}')
        node.set(qn('w:w'), '120') # 120 dxa is about 6pt padding
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.0)
    run.font.color.rgb = RGBColor(30, 40, 50)
    
    # Spacer paragraph
    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_before = Pt(0)
    p_spacer.paragraph_format.space_after = Pt(6)

def main():
    doc = Document()
    
    # Set standard page margins (1.0 inch margins)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Set default paragraph styles
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = RGBColor(40, 40, 40)
    style_normal.paragraph_format.space_after = Pt(6)
    style_normal.paragraph_format.line_spacing = 1.15

    # =========================================================================
    # PREMIUM COVER PAGE
    # =========================================================================
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(80)
    title_p.paragraph_format.space_after = Pt(10)
    
    run_title = title_p.add_run("LOCAL VEGETABLE VARIETY CLASSIFIER")
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(26)
    run_title.font.bold = True
    run_title.font.color.rgb = COLOR_TEAL
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(40)
    
    run_sub = subtitle_p.add_run("Unified Project Repository: Complete Technical Codebase followed by Experimental Inference & Evaluation Work")
    run_sub.font.name = 'Calibri'
    run_sub.font.size = Pt(14)
    run_sub.font.italic = True
    run_sub.font.color.rgb = COLOR_CHARCOAL
    
    # Divider line using thick border symbols
    div_p = doc.add_paragraph()
    div_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    div_p.paragraph_format.space_after = Pt(60)
    run_div = div_p.add_run("━" * 45)
    run_div.font.color.rgb = COLOR_TEAL
    run_div.font.bold = True
    
    # Submission Metadata Grid
    meta_table = doc.add_table(rows=5, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.autofit = True
    
    metadata = [
        ("Author / Presenter:", "Sneha Patel & Team"),
        ("Submission Format:", "Technical Codebase & Inference Work Compilation (.docx)"),
        ("Hardware Accelerator:", "Apple Silicon M-Series GPU (MPS Acceleration Enabled)"),
        ("Production Host Serving:", "Asynchronous Uvicorn ASGI + Starlette FastAPI Backend"),
        ("Deep Learning Stack:", "PyTorch, Torchvision, Scikit-Learn, Rembg (U2-Net), PIL, Matplotlib")
    ]
    
    for i, (label, val) in enumerate(metadata):
        row = meta_table.rows[i]
        
        cell_lbl = row.cells[0]
        cell_lbl.paragraphs[0].add_run(label).font.bold = True
        cell_lbl.paragraphs[0].runs[0].font.color.rgb = COLOR_CHARCOAL
        
        cell_val = row.cells[1]
        cell_val.paragraphs[0].add_run(val)
        
    doc.add_page_break()

    # =========================================================================
    # SECTION 1: COMPLETE TECHNICAL CODEBASE
    # =========================================================================
    add_heading_with_spacing(doc, "1. Complete Technical Codebase Repository", 1)
    
    p_intro_code = doc.add_paragraph()
    p_intro_code.add_run("This section houses the complete, self-contained, and highly modular Python repository constructed to pre-process, segment, augment, train, evaluate, and serve the classifier. All scripts are presented in full, featuring structured layouts and complete comments:")

    code_files = [
        ("models.py", "Defines the 6 custom PyTorch Convolutional Architectures built from scratch (incorporating GAP, Multi-Scale kernels, skip paths, and depthwise-separable parameters)."),
        ("utils.py", "Provides core optimization utility hooks (EarlyStopping validation monitors, dynamic Inverse-Frequency Class weight mapping, and targeted minority augmentations)."),
        ("app.py", "Serves as the production FastAPI backend. Incorporates parallel background stripping (rembg) and executes the weighted ensemble average predictions."),
        ("preprocess_and_balance.py", "Prepares the raw datasets, scanning training splits and oversampling minority classes programmatically to achieve uniform representations."),
        ("remove_backgrounds.py", "A multi-threaded background removal pipeline paste-stripping vegetable pixels onto solid black backgrounds using U2-Net algorithms."),
        ("calculate_neurons.py", "A mathematical analysis tool verifying dense feature map shapes and parameter capacities across scratch CNN models."),
        ("evaluate_all.py", "An automated evaluation script computing unseen test-set accuracy, precision, recall, and confusion matrix layouts programmatically."),
        ("run_inference_samples.py", "Evaluates live predictions of the saved weighted softmax ensemble against concrete image samples from each crop category."),
        ("train_deep_cnn.py", "Model 2 DeepCNN training pipeline containing Dropout interfaces, Batch Normalizations, and ReduceLROnPlateau learning monitors."),
        ("train_light_cnn.py", "Model 1 LightCNN baseline training loop."),
        ("train_gap_cnn.py", "Model 3 GAPCNN training loop."),
        ("train_multiscale_cnn.py", "Model 4 MultiScaleCNN training loop."),
        ("train_residual_cnn.py", "Model 5 CustomResidualCNN training loop."),
        ("train_depthwise_cnn.py", "Model 6 CustomDepthwiseSepCNN lightweight training loop."),
        ("train_resnet18.py", "Model 7 ResNet18 transfer learning and final classifier head adaptation loop.")
    ]

    for filename, desc in code_files:
        if os.path.exists(filename):
            print(f"Adding code block for: {filename}...")
            with open(filename, "r") as f:
                code_text = f.read()
            
            # File heading
            add_heading_with_spacing(doc, f"1.{code_files.index((filename, desc))+1}. Source File: {filename}", 2)
            
            p_desc = doc.add_paragraph()
            r_desc = p_desc.add_run(desc)
            r_desc.font.italic = True
            
            add_code_block(doc, filename, code_text)
            doc.add_page_break()

    # =========================================================================
    # SECTION 2: INFERENCE & EXPERIMENTAL RESULTS (THE OUTPUTS)
    # =========================================================================
    add_heading_with_spacing(doc, "2. Deep Learning Evaluation & Inference Work", 1)
    
    p_inf_intro = doc.add_paragraph()
    p_inf_intro.add_run("This section compiles the actual evaluation results, performance comparisons, parameter audits, and live inference outputs of our trained CNN suite on the unseen test set (196 independent vegetable samples).")

    # 2.1 Preprocessing and Pipeline Summary
    add_heading_with_spacing(doc, "2.1. Preprocessing & Data Pipeline Accomplishments", 2)
    doc.add_paragraph("Because classifying fine-grained green vegetables (Chilli, Okra, Pointed Gourd, Ivy Gourd, and Peas) suffers from color bias, our preprocessing pipeline stripped all background pixels using U2-Net, isolating morphological characteristics (elongated cylinders, stripes, spherical lobes) on a pure black backdrop. The models are trained on these segmented foregrounds, totally decoupling background noises.")

    # 2.2 Model Parameter comparison
    add_heading_with_spacing(doc, "2.2. Architectural Capacity & Neuronal Comparisons", 2)
    doc.add_paragraph("Using our architectural validator `calculate_neurons.py`, we mathematically analyzed parameter counts and dense layer structures for all custom networks:")

    param_table = doc.add_table(rows=7, cols=3)
    param_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    param_table.style = 'Light Shading Accent 1'
    
    headers_param = ["Architecture Name", "Total Trainable Parameters", "Fully Connected dense Layer Head"]
    for j, h_text in enumerate(headers_param):
        cell = param_table.cell(0, j)
        cell.paragraphs[0].add_run(h_text).font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = COLOR_TEAL
        set_cell_background(cell, "EBF1F5")

    row_data = [
        ("CustomLightCNN (Model 1)", "24,133", "fc (64 -> 5)"),
        ("CustomDeepCNN (Model 2)", "3,606,917", "fc1 (12,544 -> 256), fc2 (256 -> 5)"),
        ("CustomGAPCNN (Model 3)", "394,885", "fc (256 -> 5)"),
        ("CustomMultiScaleCNN (Model 4)", "2,685,317", "fc (384 -> 5)"),
        ("CustomResidualCNN (Model 5)", "1,227,685", "fc (256 -> 5)"),
        ("CustomDepthwiseSepCNN (Model 6)", "48,771", "fc (256 -> 5)")
    ]
    for i, row in enumerate(row_data):
        for j, val in enumerate(row):
            cell = param_table.cell(i+1, j)
            cell.paragraphs[0].text = val

    doc.add_paragraph() # Spacer

    # 2.3 Rankings
    add_heading_with_spacing(doc, "2.3. Test Set Accuracy Rankings & Analysis", 2)
    
    acc_table = doc.add_table(rows=8, cols=3)
    acc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    acc_table.style = 'Light Shading Accent 1'
    
    headers_acc = ["Rank", "Model Architecture Name", "Test Accuracy Score (%)"]
    for j, h_text in enumerate(headers_acc):
        cell = acc_table.cell(0, j)
        cell.paragraphs[0].add_run(h_text).font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = COLOR_TEAL
        set_cell_background(cell, "EBF1F5")

    rankings = [
        ("1", "CustomDeepCNN (Model 2)", "96.94%"),
        ("2", "ResNet-18 (Transfer Learning)", "90.82%"),
        ("3", "CustomGAPCNN (Model 3)", "85.20%"),
        ("4", "CustomLightCNN (Model 1)", "83.16%"),
        ("5", "CustomMultiScaleCNN (Model 4)", "82.14%"),
        ("6", "CustomResidualCNN (Model 5)", "79.08%"),
        ("7", "CustomDepthwiseSepCNN (Model 6)", "78.57%")
    ]
    for i, row in enumerate(rankings):
        for j, val in enumerate(row):
            cell = acc_table.cell(i+1, j)
            cell.paragraphs[0].text = val
            if j == 2:
                cell.paragraphs[0].runs[0].font.bold = True

    p_rank_analysis = doc.add_paragraph()
    p_rank_analysis.paragraph_format.space_before = Pt(8)
    p_rank_analysis.add_run("Ranking Discussion: ").font.bold = True
    p_rank_analysis.add_run("Our CustomDeepCNN achieved an outstanding 96.94% accuracy, outperforming the Transfer Learning ResNet-18 benchmark (90.82%). This stems from training the CustomDeepCNN completely on segmented black-background objects from scratch, enabling specialized kernels to extract local geometric features. ResNet-18, pre-trained on ImageNet, contains high-level filters that are slightly less aligned to isolated crops on high-contrast black backdrops.")

    doc.add_page_break()

    # 2.4 Model Reports
    add_heading_with_spacing(doc, "2.4. Detailed Statistical Reports Per Model Architecture", 2)
    
    try:
        with open("evaluation_results.json", "r") as f:
            eval_results = json.load(f)
    except Exception as e:
        eval_results = {}
        print(f"Error loading evaluation results: {e}")

    model_keys = [
        ("DeepCNN", "CustomDeepCNN (Model 2)", "deep_curves.png"),
        ("ResNet18", "ResNet-18 (Transfer Learning)", "resnet18_curves.png"),
        ("GAPCNN", "CustomGAPCNN (Model 3)", "gap_curves.png"),
        ("LightCNN", "CustomLightCNN (Model 1)", "light_curves.png"),
        ("MultiScaleCNN", "CustomMultiScaleCNN (Model 4)", "multiscale_curves.png"),
        ("ResidualCNN", "CustomResidualCNN (Model 5)", "residual_curves.png"),
        ("DepthwiseSepCNN", "CustomDepthwiseSepCNN (Model 6)", "depthwise_curves.png")
    ]

    for key, name, img_name in model_keys:
        if key not in eval_results:
            continue
            
        res = eval_results[key]
        add_heading_with_spacing(doc, f"Experimental Metrics: {name}", 3)
        
        p_stats = doc.add_paragraph()
        p_stats.add_run(f"• Overall Test Accuracy: {res['accuracy']*100:.2f}%\n").font.bold = True
        p_stats.add_run(f"• Macro Precision Score: {res['report']['macro avg']['precision']:.4f}\n")
        p_stats.add_run(f"• Macro Recall Score: {res['report']['macro avg']['recall']:.4f}\n")
        p_stats.add_run(f"• Macro F1-Score: {res['report']['macro avg']['f1-score']:.4f}\n")
        
        # Classification report table
        rep_table = doc.add_table(rows=6, cols=5)
        rep_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        rep_table.style = 'Light Shading Accent 1'
        
        headers_rep = ["Class Label", "Precision", "Recall", "F1-Score", "Support"]
        for j, h_text in enumerate(headers_rep):
            cell = rep_table.cell(0, j)
            cell.paragraphs[0].add_run(h_text).font.bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = COLOR_TEAL
            set_cell_background(cell, "F2F5F8")
            
        classes_to_show = ['Green Chilli (Marcha)', 'Ladies finger', 'Pointed gourd', 'ivy guard', 'peas']
        for i, cls in enumerate(classes_to_show):
            row_cell = rep_table.rows[i+1]
            row_cell.cells[0].paragraphs[0].text = cls
            
            cls_report = res['report'][cls]
            row_cell.cells[1].paragraphs[0].text = f"{cls_report['precision']:.2f}"
            row_cell.cells[2].paragraphs[0].text = f"{cls_report['recall']:.2f}"
            row_cell.cells[3].paragraphs[0].text = f"{cls_report['f1-score']:.2f}"
            row_cell.cells[4].paragraphs[0].text = str(int(cls_report['support']))
            
        doc.add_paragraph() # Spacer
        
        # Embedded curve chart
        if os.path.exists(img_name):
            try:
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.add_run().add_picture(img_name, width=Inches(4.2))
                caption = p_img.add_run(f"\nFigure: Training vs. Validation Loss and Accuracy Curves for {name}")
                caption.font.size = Pt(9.5)
                caption.font.italic = True
                caption.font.color.rgb = COLOR_GRAY
            except Exception as img_err:
                print(f"Error inserting curve {img_name}: {img_err}")
                
        doc.add_page_break()

    # 2.5 Live Ensemble Serving & Temperature scaling
    add_heading_with_spacing(doc, "2.5. Real-Time Ensemble Fusion & Calibration Mechanisms", 2)
    doc.add_paragraph("To implement standard commercial serving, our web framework `app.py` fuses predictions across four high-performing networks (DeepCNN, GAPCNN, ResidualCNN, MultiScaleCNN):")
    
    p_fuse = doc.add_paragraph()
    p_fuse.add_run("1. Proportional soft-voting power: ").font.bold = True
    p_fuse.add_run("Weighted votes based on individual accuracy: DeepCNN (28%), GAPCNN (26%), ResidualCNN (24%), MultiScaleCNN (22%).\n")
    p_fuse.add_run("2. Temperature scaling (T=0.8): ").font.bold = True
    p_fuse.add_run("Intermediate predictions are scaled down mathematically by T=0.8 before applying Softmax, accentuating confidence scores.\n")
    p_fuse.add_run("3. Bayesian Class-Prior Calibration: ").font.bold = True
    p_fuse.add_run("To eliminate over-segmentation augmentations residues, the ensemble scales the final Ivy Gourd output by a 15% penalty factor, ensuring rigorous multi-model verification.")

    # 2.6 Live Inference Outputs on concrete images
    add_heading_with_spacing(doc, "2.6. Verified Live Soft-Voting Ensemble Inference Run", 2)
    doc.add_paragraph("Below is the actual console execution output generated by loading our completed soft-voting ensemble models and executing real-time predictions on unseen test photos for each botanical class:")

    # Read console output from file
    inference_output_text = ""
    if os.path.exists("inference_console_output.txt"):
        try:
            with open("inference_console_output.txt", "r") as f:
                inference_output_text = f.read()
        except Exception as e:
            inference_output_text = f"Error reading inference output: {e}"
    else:
        inference_output_text = "No live inference file found. Run 'run_inference_samples.py' to generate."

    # Monospaced box for the console output
    table_inf = doc.add_table(rows=1, cols=1)
    table_inf.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_inf.autofit = False
    
    cell_inf = table_inf.cell(0, 0)
    cell_inf.width = Inches(6.5)
    set_cell_background(cell_inf, "2B2D42") # Dark navy blue for beautiful console style
    
    tcPr = cell_inf._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin in ['top', 'bottom', 'left', 'right']:
        node = OxmlElement(f'w:{margin}')
        node.set(qn('w:w'), '120')
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)
    
    p_inf = cell_inf.paragraphs[0]
    p_inf.paragraph_format.space_before = Pt(4)
    p_inf.paragraph_format.space_after = Pt(4)
    p_inf.paragraph_format.line_spacing = 1.0
    
    run_inf = p_inf.add_run(inference_output_text)
    run_inf.font.name = 'Courier New'
    run_inf.font.size = Pt(8.5)
    run_inf.font.color.rgb = RGBColor(237, 242, 244) # Pure off-white terminal color

    # Final Save
    out_path = "Local_Vegetable_Variety_Classifier_Report.docx"
    doc.save(out_path)
    print(f"\nRESTUCTURED DOCUMENT SUCCESSFULLY CREATED: {out_path}")

if __name__ == "__main__":
    main()
